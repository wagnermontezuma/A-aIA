# memory/document_processor.py
import asyncio
import aiohttp
import os
import tempfile
import hashlib
from typing import List, Dict, Optional, Union
from pathlib import Path
import mimetypes
from datetime import datetime
from dotenv import load_dotenv

# Processamento de diferentes tipos de arquivo
try:
    import PyPDF2
    import pdfplumber
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from bs4 import BeautifulSoup
    import requests
    WEB_SCRAPING_AVAILABLE = True
except ImportError:
    WEB_SCRAPING_AVAILABLE = False

try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

class DocumentProcessor:
    """
    Processador de documentos para alimentação do RAG.
    Suporta URLs, PDFs, TXT, DOCX e outros formatos.
    """
    
    def __init__(self):
        load_dotenv()
        self.supported_formats = {
            '.txt': self._process_text,
            '.md': self._process_text,
            '.pdf': self._process_pdf,
            '.docx': self._process_docx,
            '.html': self._process_html,
            '.htm': self._process_html
        }
        
        self.max_file_size = 50 * 1024 * 1024  # 50MB
        self.chunk_size = 1000  # Tamanho dos chunks
        self.chunk_overlap = 200  # Sobreposição entre chunks
    
    async def process_url(self, url: str, source_name: str = None) -> Dict:
        """Processa uma URL e extrai o conteúdo."""
        if not WEB_SCRAPING_AVAILABLE:
            raise ImportError("beautifulsoup4 e requests são necessários para processar URLs")
        
        try:
            print(f"🌐 Processando URL: {url}")
            
            async with aiohttp.ClientSession() as session:
                async with session.get(url, timeout=30) as response:
                    if response.status != 200:
                        raise Exception(f"Erro HTTP {response.status}")
                    
                    content_type = response.headers.get('content-type', '').lower()
                    content = await response.text()
            
            # Detectar tipo de conteúdo
            if 'pdf' in content_type:
                # Para PDFs via URL, baixar primeiro
                async with aiohttp.ClientSession() as session:
                    async with session.get(url) as response:
                        pdf_content = await response.read()
                        return await self._process_pdf_bytes(pdf_content, source_name or url)
            
            elif 'html' in content_type or url.startswith('http'):
                return await self._process_html_content(content, source_name or url)
            
            else:
                # Tratar como texto
                return await self._process_text_content(content, source_name or url)
            
        except Exception as e:
            print(f"❌ Erro ao processar URL {url}: {e}")
            raise
    
    async def process_file(self, file_path: str, source_name: str = None) -> Dict:
        """Processa um arquivo local."""
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"Arquivo não encontrado: {file_path}")
        
        if file_path.stat().st_size > self.max_file_size:
            raise ValueError(f"Arquivo muito grande: {file_path.stat().st_size} bytes")
        
        extension = file_path.suffix.lower()
        
        if extension not in self.supported_formats:
            raise ValueError(f"Formato não suportado: {extension}")
        
        print(f"📄 Processando arquivo: {file_path.name}")
        
        processor = self.supported_formats[extension]
        return await processor(file_path, source_name or file_path.name)
    
    async def process_discord_attachment(self, attachment_url: str, filename: str, source_name: str = None) -> Dict:
        """Processa um anexo do Discord."""
        try:
            print(f"📎 Processando anexo Discord: {filename}")
            
            # Baixar o arquivo
            async with aiohttp.ClientSession() as session:
                async with session.get(attachment_url) as response:
                    if response.status != 200:
                        raise Exception(f"Erro ao baixar anexo: HTTP {response.status}")
                    
                    content = await response.read()
            
            # Detectar tipo de arquivo
            extension = Path(filename).suffix.lower()
            
            if extension == '.pdf' and PDF_AVAILABLE:
                return await self._process_pdf_bytes(content, source_name or filename)
            
            elif extension in ['.txt', '.md']:
                text_content = content.decode('utf-8', errors='ignore')
                return await self._process_text_content(text_content, source_name or filename)
            
            elif extension in ['.html', '.htm']:
                html_content = content.decode('utf-8', errors='ignore')
                return await self._process_html_content(html_content, source_name or filename)
            
            else:
                raise ValueError(f"Formato de anexo não suportado: {extension}")
            
        except Exception as e:
            print(f"❌ Erro ao processar anexo {filename}: {e}")
            raise
    
    async def _process_text(self, file_path: Path, source_name: str) -> Dict:
        """Processa arquivos de texto."""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()
        
        return await self._process_text_content(content, source_name)
    
    async def _process_text_content(self, content: str, source_name: str) -> Dict:
        """Processa conteúdo de texto."""
        chunks = self._split_text(content)
        
        return {
            'source': source_name,
            'type': 'text',
            'total_chunks': len(chunks),
            'chunks': chunks,
            'metadata': {
                'processed_at': datetime.now().isoformat(),
                'char_count': len(content),
                'word_count': len(content.split())
            }
        }
    
    async def _process_pdf(self, file_path: Path, source_name: str) -> Dict:
        """Processa arquivos PDF."""
        if not PDF_AVAILABLE:
            raise ImportError("PyPDF2 ou pdfplumber são necessários para processar PDFs")
        
        with open(file_path, 'rb') as f:
            content = f.read()
        
        return await self._process_pdf_bytes(content, source_name)
    
    async def _process_pdf_bytes(self, pdf_bytes: bytes, source_name: str) -> Dict:
        """Processa bytes de PDF."""
        try:
            # Tentar com pdfplumber primeiro (melhor para texto)
            import io
            text_content = ""
            
            if 'pdfplumber' in globals():
                import pdfplumber
                with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text_content += page_text + "\n"
            
            # Fallback para PyPDF2
            if not text_content and PDF_AVAILABLE:
                import PyPDF2
                pdf_reader = PyPDF2.PdfReader(io.BytesIO(pdf_bytes))
                for page in pdf_reader.pages:
                    text_content += page.extract_text() + "\n"
            
            if not text_content:
                raise Exception("Não foi possível extrair texto do PDF")
            
            chunks = self._split_text(text_content)
            
            return {
                'source': source_name,
                'type': 'pdf',
                'total_chunks': len(chunks),
                'chunks': chunks,
                'metadata': {
                    'processed_at': datetime.now().isoformat(),
                    'char_count': len(text_content),
                    'word_count': len(text_content.split()),
                    'pages': len(pdf_reader.pages) if 'pdf_reader' in locals() else 'unknown'
                }
            }
            
        except Exception as e:
            print(f"❌ Erro ao processar PDF: {e}")
            raise
    
    async def _process_html(self, file_path: Path, source_name: str) -> Dict:
        """Processa arquivos HTML."""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()
        
        return await self._process_html_content(content, source_name)
    
    async def _process_html_content(self, html_content: str, source_name: str) -> Dict:
        """Processa conteúdo HTML."""
        if not WEB_SCRAPING_AVAILABLE:
            raise ImportError("beautifulsoup4 é necessário para processar HTML")
        
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # Remover scripts e estilos
        for script in soup(["script", "style"]):
            script.decompose()
        
        # Extrair texto
        text_content = soup.get_text()
        
        # Limpar texto
        lines = (line.strip() for line in text_content.splitlines())
        chunks_text = '\n'.join(chunk for chunk in lines if chunk)
        
        chunks = self._split_text(chunks_text)
        
        return {
            'source': source_name,
            'type': 'html',
            'total_chunks': len(chunks),
            'chunks': chunks,
            'metadata': {
                'processed_at': datetime.now().isoformat(),
                'char_count': len(chunks_text),
                'word_count': len(chunks_text.split()),
                'title': soup.title.string if soup.title else None
            }
        }
    
    async def _process_docx(self, file_path: Path, source_name: str) -> Dict:
        """Processa arquivos DOCX."""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx é necessário para processar arquivos DOCX")
        
        import docx
        doc = docx.Document(file_path)
        
        text_content = ""
        for paragraph in doc.paragraphs:
            text_content += paragraph.text + "\n"
        
        chunks = self._split_text(text_content)
        
        return {
            'source': source_name,
            'type': 'docx',
            'total_chunks': len(chunks),
            'chunks': chunks,
            'metadata': {
                'processed_at': datetime.now().isoformat(),
                'char_count': len(text_content),
                'word_count': len(text_content.split()),
                'paragraphs': len(doc.paragraphs)
            }
        }
    
    def _split_text(self, text: str) -> List[str]:
        """Divide o texto em chunks com sobreposição."""
        if len(text) < self.chunk_size:
            return [text]
        
        chunks = []
        start = 0
        while start < len(text):
            end = start + self.chunk_size
            chunk = text[start:end]
            chunks.append(chunk)
            start = end - self.chunk_overlap if end < len(text) else end
            
            if start >= len(text):
                break
        
        return chunks